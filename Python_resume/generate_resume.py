"""
Edit this file to update your resume, then run:
    python generate_resume.py

It will regenerate: Gautam_Kumar_Resume.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


# =====================
# EDIT YOUR DATA HERE
# =====================

NAME = "GAUTAM KUMAR"
ROLE = "Web Developer | Python Learner | Linux & Networking Enthusiast"

EMAIL = "your@email.com"
GITHUB = "github.com/yourusername"
LINKEDIN = "linkedin.com/in/yourusername"

PROFILE = (
    "Aspiring Web Developer with hands-on experience in HTML, CSS, and JavaScript, "
    "currently learning backend development with Python. Strong foundation in Linux "
    "and Networking, with an interest in building secure and efficient applications. "
    "Focused on practical learning through real-world projects."
)

SKILLS = [
    "HTML, CSS, JavaScript",
    "Responsive Web Design",
    "Python (Tkinter, Basics)",
    "Linux Command Line",
    "Networking Fundamentals",
    "Git & GitHub",
    "Problem Solving",
    "Security Awareness",
]

PROJECTS = [
    (
        "Personal Portfolio Website",
        "Developed a responsive portfolio using HTML, CSS, and JavaScript with structured layout and modern UI."
    ),
    (
        "Alarm Clock Application (Python – Tkinter)",
        "Created a desktop alarm clock using Python Tkinter with GUI interface, real-time clock display, and alert scheduling."
    ),
]

CERTIFICATIONS = [
    "Python Basic",
    "Frontend Developer – Coding Ninjas",
    "Java – Coding Ninjas",
    "Backend Developer – Coding Ninjas",
    "Networking – Cisco",
]

EDUCATION = "Your Course / Degree — Your Institution (Edit This)"

OUTPUT_FILE = "Gautam_Kumar_Resume.docx"

# =====================
# DO NOT EDIT BELOW
# =====================

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

# Name
p = doc.add_paragraph()
run = p.add_run(NAME)
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)   # Blue color
run.font.size = Pt(18)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Role
p = doc.add_paragraph()
run = p.add_run(ROLE)
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Contact
p = doc.add_paragraph()
run = p.add_run(f"{EMAIL} | {GITHUB} | {LINKEDIN}")
run.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Profile
p = doc.add_paragraph()
run = p.add_run("PROFILE")
run.bold = True
doc.add_paragraph(PROFILE)

# Skills
p = doc.add_paragraph()
run = p.add_run("\\nTECHNICAL SKILLS")
run.bold = True

for skill in SKILLS:
    doc.add_paragraph(skill).style = doc.styles["List Bullet"]

# Projects
p = doc.add_paragraph()
run = p.add_run("\\nPROJECTS")
run.bold = True

for title, desc in PROJECTS:
    doc.add_paragraph(title).runs[0].bold = True
    doc.add_paragraph(desc)

# Certifications
p = doc.add_paragraph()
run = p.add_run("\\nCERTIFICATIONS")
run.bold = True

for cert in CERTIFICATIONS:
    doc.add_paragraph(cert).style = doc.styles["List Bullet"]

# Education
p = doc.add_paragraph()
run = p.add_run("\\nEDUCATION")
run.bold = True
doc.add_paragraph(EDUCATION)

doc.save(OUTPUT_FILE)

print(f"Resume generated successfully: {OUTPUT_FILE}")
